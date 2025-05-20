# from fastapi import FastAPI, UploadFile, File, Form
# from fastapi.responses import JSONResponse, StreamingResponse
# from fastapi.middleware.cors import CORSMiddleware
# from transformers import TrOCRProcessor, VisionEncoderDecoderModel
# from PIL import Image
# import io
# from docx import Document
# from reportlab.pdfgen import canvas
# import traceback

# app = FastAPI()

# # Allow all CORS for development purposes (limit this in production!)
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # Load TrOCR model and processor
# processor = TrOCRProcessor.from_pretrained("microsoft/trocr-base-handwritten")
# model = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-base-handwritten")


# @app.post("/convert")
# async def convert_image(file: UploadFile = File(...)):
#     try:
#         print("Received file:", file.filename)
#         image_bytes = await file.read()
#         image = Image.open(io.BytesIO(image_bytes)).convert("RGB")

#         pixel_values = processor(images=image, return_tensors="pt").pixel_values
#         generated_ids = model.generate(pixel_values)
#         text = processor.batch_decode(generated_ids, skip_special_tokens=True)[0]

#         print("OCR Result:", text)
#         return {"text": text}

#     except Exception as e:
#         print("=== OCR FAILED ===")
#         print(traceback.format_exc())
#         return JSONResponse(status_code=500, content={"error": str(e)})


# @app.post("/save")
# async def save_file(
#     text: str = Form(...),
#     format: str = Form(...)
# ):
#     try:
#         buffer = io.BytesIO()

#         if format.lower() == "pdf":
#             p = canvas.Canvas(buffer)
#             y = 800
#             for line in text.splitlines():
#                 p.drawString(40, y, line)
#                 y -= 15
#             p.save()
#             buffer.seek(0)
#             return StreamingResponse(buffer, media_type="application/pdf", headers={
#                 "Content-Disposition": "attachment; filename=output.pdf"
#             })

#         elif format.lower() == "docx":
#             doc = Document()
#             for line in text.splitlines():
#                 doc.add_paragraph(line)
#             doc.save(buffer)
#             buffer.seek(0)
#             return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={
#                 "Content-Disposition": "attachment; filename=output.docx"
#             })

#         else:
#             return JSONResponse(status_code=400, content={"error": "Unsupported format"})

#     except Exception as e:
#         return JSONResponse(status_code=500, content={"error": "File save failed", "details": str(e)})



#Previous working version:
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
from PIL import Image, ImageOps, ImageEnhance
import io
from docx import Document
from reportlab.pdfgen import canvas
import traceback

app = FastAPI()

#CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load TrOCR model and processor
processor = TrOCRProcessor.from_pretrained("microsoft/trocr-large-handwritten")
model = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-large-handwritten")

def preprocess_image(image: Image.Image) -> Image.Image:
    """Apply enhancements to improve OCR accuracy"""
    image = image.convert("L")  # Grayscale
    image = ImageOps.invert(image)  # Invert for white background
    image = ImageOps.autocontrast(image)  # Enhance contrast
    image = image.resize((1024, 1280), Image.Resampling.LANCZOS)  # Resize using proper method
    image = ImageEnhance.Sharpness(image).enhance(2.0)  # Sharpen text
    return image.convert("RGB")  # Return RGB image

@app.post("/convert")
async def convert_image(file: UploadFile = File(...)):
    try:
        print("Received file:", file.filename)
        image_bytes = await file.read()
        image = Image.open(io.BytesIO(image_bytes)).convert("RGB")

        #Preprocessing the image
        preprocessed_image = preprocess_image(image)

        #OCR
        pixel_values = processor(images=preprocessed_image, return_tensors="pt").pixel_values
        generated_ids = model.generate(pixel_values)
        text = processor.batch_decode(generated_ids, skip_special_tokens=True)[0]

        print("OCR Result:", text)
        return {"text": text}

    except Exception as e:
        print("=== OCR FAILED ===")
        print(traceback.format_exc())
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/save")
async def save_file(
    text: str = Form(...),
    format: str = Form(...)
):
    try:
        buffer = io.BytesIO()

        if format.lower() == "pdf":
            p = canvas.Canvas(buffer)
            p.setFont("Helvetica", 20)
            y = 800
            for line in text.splitlines():
                p.drawString(40, y, line)
                y -= 15
            p.save()
            buffer.seek(0)
            return StreamingResponse(buffer, media_type="application/pdf", headers={
                "Content-Disposition": "attachment; filename=output.pdf"
            })

        elif format.lower() == "docx":
            doc = Document()
            for line in text.splitlines():
                doc.add_paragraph(line)
            doc.save(buffer)
            buffer.seek(0)
            return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={
                "Content-Disposition": "attachment; filename=output.docx"
            })

        else:
            return JSONResponse(status_code=400, content={"error": "Unsupported format"})

    except Exception as e:
        return JSONResponse(status_code=500, content={"error": "File save failed", "details": str(e)})


# from fastapi import FastAPI, UploadFile, File, Form
# from fastapi.responses import JSONResponse, StreamingResponse
# from fastapi.middleware.cors import CORSMiddleware
# from PIL import Image
# import io
# from docx import Document
# from reportlab.pdfgen import canvas
# import traceback
# import easyocr

# app = FastAPI()

# # CORS for development (relax in prod)
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # Load EasyOCR reader (Tamil + English)
# reader = easyocr.Reader(['ta', 'en'], gpu=False)

# @app.post("/convert")
# async def convert_image(file: UploadFile = File(...)):
#     try:
#         print("Received file:", file.filename)
#         image_bytes = await file.read()
#         image = Image.open(io.BytesIO(image_bytes)).convert("RGB")

#         # Perform OCR with EasyOCR
#         results = reader.readtext(image_bytes, detail=0)
#         text = "\n".join(results)
#         print("OCR Result:", text)

#         return {"text": text}

#     except Exception as e:
#         print("=== OCR FAILED ===")
#         print(traceback.format_exc())
#         return JSONResponse(status_code=500, content={"error": str(e)})

# @app.post("/save")
# async def save_file(
#     text: str = Form(...),
#     format: str = Form(...)
# ):
#     try:
#         buffer = io.BytesIO()

#         if format.lower() == "pdf":
#             p = canvas.Canvas(buffer)
#             y = 800
#             for line in text.splitlines():
#                 p.drawString(40, y, line)
#                 y -= 15
#             p.save()
#             buffer.seek(0)
#             return StreamingResponse(buffer, media_type="application/pdf", headers={
#                 "Content-Disposition": "attachment; filename=output.pdf"
#             })

#         elif format.lower() == "docx":
#             doc = Document()
#             for line in text.splitlines():
#                 doc.add_paragraph(line)
#             doc.save(buffer)
#             buffer.seek(0)
#             return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={
#                 "Content-Disposition": "attachment; filename=output.docx"
#             })

#         else:
#             return JSONResponse(status_code=400, content={"error": "Unsupported format"})

#     except Exception as e:
#         return JSONResponse(status_code=500, content={"error": "File save failed", "details": str(e)})

